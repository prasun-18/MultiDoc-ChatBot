import os
os.environ['TF_CPP_MIN_LOG_LEVEL'] = '3'  # Suppress TensorFlow logs
os.environ['CUDA_VISIBLE_DEVICES'] = '-1'  # Disable GPU

import PyPDF2
import docx
import pandas as pd
import json
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np
from transformers import pipeline
import streamlit as st
import base64
from docx import Document
import time

# Load a pre-trained sentence embedding model
embedding_model = SentenceTransformer('all-MiniLM-L6-v2')

# Load a Hugging Face LLM for question answering
qa_pipeline = pipeline("question-answering", model="deepset/roberta-base-squad2")

# Initialize FAISS index for vector storage
dimension = 384  # Dimension of the embeddings
index = faiss.IndexFlatL2(dimension)

# Dictionary to store text chunks and their metadata
text_data = []

def read_pdf(file_path):
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ''
        for page in reader.pages:
            text += page.extract_text()
        return text

def read_txt(file_path):
    with open(file_path, 'r') as file:
        return file.read()

def read_docx(file_path):
    doc = docx.Document(file_path)
    text = ''
    for para in doc.paragraphs:
        text += para.text + '\n'
    return text

def read_excel(file_path):
    df = pd.read_excel(file_path)
    return df.to_string()

def read_csv(file_path):
    df = pd.read_csv(file_path)
    return df.to_string()

def read_json(file_path):
    with open(file_path, 'r') as file:
        data = json.load(file)
    return json.dumps(data)

def read_file(file_path):
    if file_path.endswith('.pdf'):
        return read_pdf(file_path)
    elif file_path.endswith('.txt'):
        return read_txt(file_path)
    elif file_path.endswith('.docx'):
        return read_docx(file_path)
    elif file_path.endswith('.xlsx') or file_path.endswith('.xls'):
        return read_excel(file_path)
    elif file_path.endswith('.csv'):
        return read_csv(file_path)
    elif file_path.endswith('.json'):
        return read_json(file_path)
    else:
        raise ValueError("Unsupported file format")

def chunk_text(text, chunk_size=500):
    """Split text into smaller chunks for embedding."""
    words = text.split()
    chunks = [' '.join(words[i:i + chunk_size]) for i in range(0, len(words), chunk_size)]
    return chunks

def add_to_vector_db(text_chunks):
    """Generate embeddings and add them to the FAISS index."""
    global text_data, index
    embeddings = embedding_model.encode(text_chunks)
    for i, (chunk, embedding) in enumerate(zip(text_chunks, embeddings)):
        text_data.append({"id": len(text_data), "text": chunk})
        index.add(np.array([embedding]))

def query_vector_db(query, top_k=3):
    """Search the vector database for the most relevant chunks."""
    query_embedding = embedding_model.encode([query])
    distances, indices = index.search(np.array(query_embedding), top_k)
    results = []
    for i, idx in enumerate(indices[0]):
        if idx >= 0:
            results.append(text_data[idx]["text"])
    return results

def answer_question(query):
    """Answer the user's question using the vector database and LLM."""
    # Search the vector database for relevant chunks
    relevant_chunks = query_vector_db(query)
    
    # If no relevant chunks are found, search external sources
    if not relevant_chunks:
        return "No answer found in the uploaded files. Please check external sources."
    
    # Use the LLM to generate answers from the relevant chunks
    answers = []
    for chunk in relevant_chunks:
        result = qa_pipeline(question=query, context=chunk)
        answers.append(result["answer"])
    
    # Return the top 3 answers
    return answers[:3]

# Sidebar
st.sidebar.title("File Space")

# Function to show PDF in the sidebar
def show_pdf(pdf_file, width=300, height=200):
    pdf_data = pdf_file.read()
    pdf_base64 = base64.b64encode(pdf_data).decode("utf-8")
    pdf_embed_code = f'<iframe src="data:application/pdf;base64,{pdf_base64}" width="{width}" height="{height}" type="application/pdf"></iframe>'
    st.sidebar.markdown(pdf_embed_code, unsafe_allow_html=True)

# Function to show text files in the sidebar
def show_text(file, width=300, height=200):
    content = file.read().decode("utf-8")
    html_content = f'<textarea style="width:{width}px;height:{height}px;">{content[:1000]}</textarea>'
    st.sidebar.markdown(html_content, unsafe_allow_html=True)
    return content

# Function to show JSON files in the sidebar
def show_json(file, width=300, height=200):
    content = json.load(file)
    json_str = json.dumps(content, indent=4)
    html_content = f'<textarea style="width:{width}px;height:{height}px;">{json_str[:2000]}</textarea>'
    st.sidebar.markdown(html_content, unsafe_allow_html=True)
    return content

# Function to show data files (CSV/Excel) in the sidebar
def show_data(file, file_type, width=300, height=200):
    if file_type == "csv":
        df = pd.read_csv(file)
    elif file_type == "xlsx":
        df = pd.read_excel(file)
    st.sidebar.dataframe(df.head(), width=width, height=height)
    return df

# Function to show Word documents in the sidebar
def show_word(file, width=300, height=200):
    doc = Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    html_content = f'<textarea style="width:{width}px;height:{height}px;">{text[:1000]}</textarea>'
    st.sidebar.markdown(html_content, unsafe_allow_html=True)
    return text

# Sidebar file uploader
uploaded_file = st.sidebar.file_uploader(
    "Upload a file", type=["csv", "txt", "xlsx", "json", "pdf", "docx"]
)

# Display the uploaded file
if uploaded_file is not None:
    st.sidebar.success(f"File {uploaded_file.name} uploaded successfully!")
    file_path = os.path.join("/tmp", uploaded_file.name)
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    
    # Read the file and process it
    text = read_file(file_path)
    chunks = chunk_text(text)
    add_to_vector_db(chunks)

    # Determine the file type and process accordingly
    if uploaded_file.name.endswith(".csv"):
        data = show_data(uploaded_file, "csv")
    elif uploaded_file.name.endswith(".txt"):
        content = show_text(uploaded_file)
    elif uploaded_file.name.endswith(".xlsx"):
        data = show_data(uploaded_file, "xlsx")
    elif uploaded_file.name.endswith(".json"):
        content = show_json(uploaded_file)
    elif uploaded_file.name.endswith(".pdf"):
        show_pdf(uploaded_file)
    elif uploaded_file.name.endswith(".docx"):
        content = show_word(uploaded_file)
    else:
        st.sidebar.warning("Unsupported file format.")

# Main Interface
st.title("DoxBot - version1")

with st.chat_message("assistant"):
    st.write("Hello there! How can I assist you today?")

# Initialize session state for storing conversation history
if "messages" not in st.session_state:
    st.session_state.messages = []

# Display the conversation
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# Input box for user to type their message
if user_input := st.chat_input("Type your message here..."):
    # Add user message to the conversation
    st.session_state.messages.append({"role": "user", "content": user_input})

    # Display user message immediately
    with st.chat_message("user"):
        st.markdown(f"User: {user_input}")

    # Generate bot response
    bot_response = answer_question(user_input)
    if isinstance(bot_response, list):
        bot_response = "\n".join([f"{i + 1}. {answer}" for i, answer in enumerate(bot_response)])
    else:
        bot_response = f"Bot: {bot_response}"

    # Add bot message to the conversation
    st.session_state.messages.append({"role": "assistant", "content": bot_response})

    # Simulate chatbot "typing"
    with st.chat_message("assistant"):
        message_placeholder = st.empty()
        generated_text = ""
        for char in bot_response:
            generated_text += char
            message_placeholder.text(generated_text)
            time.sleep(0.03)
        message_placeholder.text(bot_response)

# Button to clear chat history
if st.button("Clear Chat", help="Click to clear the chat"):
    st.session_state.messages = []
    st.rerun()